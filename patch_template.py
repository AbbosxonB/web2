from docx import Document
import os

path = r"apps\results\templates\results\docx\vedmost_template.docx"
if not os.path.exists(path):
    print("File not found")
    exit(1)

doc = Document(path)

replacements = {
    '{Guruh}': '{{ Guruh }}',
    '{Fakultet}': '{{ Fakultet }}',
    '{Kursi}': '{{ Kursi }}',
    '{Test_nomi}': '{{ Test_nomi }}',
    '{Sana}': '{{ Sana }}',
    'Test nomi:': 'Test nomi: {{ Test_nomi }}', # In case it's empty
    # Table logic is harder.
    # The table row has: '1', 'Full_name', 'Result ', 'Null'
    # We need to replace the ROW with a loop.
}

# Simple text replacement for paragraphs
for p in doc.paragraphs:
    for key, val in replacements.items():
        if key in p.text:
            print(f"Replacing {key} in paragraph")
            p.text = p.text.replace(key, val)

# Table replacement
# We look for the table with headers.
for table in doc.tables:
    # Check headers
    if len(table.rows) > 0:
        header_cells = [c.text.strip() for c in table.rows[0].cells]
        if 'T/r' in header_cells or '№' in header_cells:
            print("Found results table")
            # We assume the second row is the template row
            if len(table.rows) > 1:
                row = table.rows[1]
                # We need to put jinja2 tags in the cells:
                # Col 0: {{ r.number }}
                # Col 1: {{ r.full_name }}
                # Col 2: {{ r.score }}
                # Col 3: {{ r.signature }}
                
                # We also need the loop tag `{% tr for r in results_table %}` but docxtpl handles rows differently usually.
                # Standard docxtpl: put `{% tr for r in results_table %}` in the first cell, and `{% tr endfor %}` in the last cell (or use merge).
                # Actually, simpler is:
                # Cell 0: {% for r in results_table %}{{ r.number }}
                # Cell 1: {{ r.full_name }}
                # Cell 2: {{ r.score }}
                # Cell last: {{ r.signature }}{% endfor %}
                
                # Check cell count
                cells = row.cells
                if len(cells) >= 3:
                     cells[0].text = "{% for r in results_table %}" + "{{ r.number }}"
                     cells[1].text = "{{ r.full_name }}"
                     cells[2].text = "{{ r.score }}"
                     if len(cells) > 3:
                         cells[3].text = "{% endfor %}"
                     else:
                         cells[2].text += "{% endfor %}"
                         
                     print("Table row patched")

doc.save(path)
print("Done")
